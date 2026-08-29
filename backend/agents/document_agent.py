import docx
import pdfplumber

class DocumentAgent:
    def analyze_docx(self, file_path: str) -> dict:
        try:
            doc = docx.Document(file_path)
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
            raw_text = "\n".join([p.text for p in doc.paragraphs])
            return {
                "type": "docx",
                "headings": headings,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "raw_text": raw_text
            }
        except Exception as e:
            return {"error": str(e)}

    def analyze_pdf(self, file_path: str) -> dict:
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
            return {
                "type": "pdf",
                "text_length": len(text),
                "raw_text": text
            }
        except Exception as e:
            return {"error": str(e)}
