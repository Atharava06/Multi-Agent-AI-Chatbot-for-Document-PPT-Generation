import docx
from services.gemini import gemini_service
import uuid
from pathlib import Path

class GenerationAgent:
    def generate(self, context: str, template_filename: str = None) -> str:
        prompt = f"Generate ONLY the long-form document/report/proposal content based on this context:\n{context}\n\nIMPORTANT: Output plain text ONLY. DO NOT use markdown bold (**), italics, or asterisks (*). If you want to create a heading, start the line exactly with 'HEADING: '. If you want to create a bullet point, start the line exactly with 'BULLET: '.\n\nCRITICAL: DO NOT generate or include any presentation slides, blueprints, or slide deck content. Ignore any context related to presentations and focus ONLY on the long-form document."
        content = gemini_service.generate_text(prompt)
        
        doc = docx.Document()
        
        if template_filename:
            template_path = Path(f"uploads/{template_filename}")
            if template_path.exists():
                try:
                    doc = docx.Document(template_path)
                    doc.add_page_break()
                except:
                    pass
                    
        doc.add_heading('AI Generated Content', 0)
        
        for p in content.split('\n'):
            p = p.strip()
            if not p:
                continue
            
            # Clean any stray markdown just in case
            import re
            p = re.sub(r'\*\*(.*?)\*\*', r'\1', p)
            p = re.sub(r'##+', '', p).strip()
            
            if p.startswith('HEADING:'):
                doc.add_heading(p.replace('HEADING:', '').strip(), level=1)
            elif p.startswith('BULLET:'):
                try:
                    doc.add_paragraph(p.replace('BULLET:', '').strip(), style='List Bullet')
                except KeyError:
                    # Fallback if List Bullet style doesn't exist in template
                    doc.add_paragraph("• " + p.replace('BULLET:', '').strip())
            else:
                doc.add_paragraph(p)
                
        filename = f"generated_{uuid.uuid4().hex[:8]}.docx"
        doc.save(f"generated/{filename}")
        
        return filename
